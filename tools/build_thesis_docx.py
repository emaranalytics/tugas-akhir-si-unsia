#!/usr/bin/env python
"""Generator naskah Tugas Akhir (.docx) — Universitas Siber Asia.

Membaca file Markdown bab di `bab/` dan merangkainya menjadi satu naskah .docx
lengkap dengan front matter (cover s.d. daftar gambar) sesuai template Lampiran
UNSIA. Bab yang belum ditulis dirender sebagai placeholder + outline sub-bab.

Cara pakai:
    python tools/build_thesis_docx.py

Untuk memperbarui setelah menyelesaikan sebuah bab: tulis/lengkapi file
markdown bab tersebut di `bab/`, daftarkan di CHAPTERS, lalu jalankan ulang.
Buka hasil di Word lalu tekan "Update Field" (Ctrl+A, F9) untuk mengisi Daftar Isi.

Format mengikuti Pedoman Teknis Penulisan TA UNSIA:
A4; margin atas/bawah/kanan 3 cm, kiri 4 cm; Times New Roman 12; spasi 1,5;
isi tabel font 10 spasi single; istilah asing miring; sitasi IEEE.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

# --------------------------------------------------------------------------- #
# Konfigurasi metadata naskah
# --------------------------------------------------------------------------- #
ROOT = Path(__file__).resolve().parent.parent
BAB_DIR = ROOT / "bab"
OUT_DIR = ROOT / "draft"
LOGO = ROOT / "reference" / "logo.png"
OUT_FILE = OUT_DIR / "Draft-Tugas-Akhir-Muhammadridwan.docx"

META = {
    "judul": "IMPLEMENTASI DAN EVALUASI TOOL REGISTRY UNTUK SKALABILITAS "
    "AI AGENT MULTI-MODUL PADA PLATFORM ERP RESTORAN ZERLO.ID",
    "nama": "Muhammadridwan",
    "nim": "220101010009",
    "prodi": "Sistem Informasi",
    "gelar": "Sarjana Komputer",
    "pembimbing": "Ikhwani Saputra, S.Kom., M.Kom.",
    "kota": "Jakarta",
    "bulan_tahun": "Juni 2026",
    "tahun": "2026",
}

# Daftar bab: (markdown_file | None, label, judul, outline_untuk_placeholder)
CHAPTERS = [
    ("bab-1-pendahuluan.md", "BAB I", "PENDAHULUAN", None),
    ("bab-2-landasan-teori.md", "BAB II", "LANDASAN TEORI", None),
    ("bab-3-implementasi-metode.md", "BAB III", "IMPLEMENTASI METODE USULAN", None),
    ("bab-4-hasil-analisa.md", "BAB IV", "HASIL DAN ANALISA", None),
    ("bab-5-kesimpulan.md", "BAB V", "KESIMPULAN", None),
]

# Daftar Pustaka konsolidasi (IEEE, urut kemunculan pertama di naskah Bab I–V).
# Sumber sitasi in-text [n] di bab/*.md sudah dinomori ulang agar konsisten
# dengan daftar ini. Format inline `*...*` dirender miring oleh add_runs().
REFERENCES = [
    'OpenAI, "Function calling and other API updates," *OpenAI Blog*, Jun. 2023. '
    "[Online]. Tersedia: https://openai.com/index/function-calling-and-other-api-updates/",
    'Google DeepMind, "Gemini API function calling," *Google AI Developer '
    "Documentation*, 2024. [Online]. Tersedia: "
    "https://ai.google.dev/gemini-api/docs/function-calling",
    "T. Schick, J. Dwivedi-Yu, R. Dessi, R. Raileanu, M. Lomeli, E. Hambro, "
    'L. Zettlemoyer, N. Cancedda, dan T. Scialom, "Toolformer: Language models '
    'can teach themselves to use tools," in *Proc. Adv. Neural Inf. Process. '
    "Syst. (NeurIPS)*, 2023. arXiv:2302.04761.",
    "S. Yao, J. Zhao, D. Yu, N. Du, I. Shafran, K. Narasimhan, dan Y. Cao, "
    '"ReAct: Synergizing reasoning and acting in language models," in *Proc. '
    "Int. Conf. Learn. Representations (ICLR)*, 2023. arXiv:2210.03629.",
    "N.F. Liu, K. Lin, J. Hewitt, A. Paranjape, M. Bevilacqua, F. Petroni, dan "
    'P. Liang, "Lost in the middle: How language models use long contexts," '
    "*Trans. Assoc. Comput. Linguistics*, vol. 12, pp. 157–173, 2024. "
    "arXiv:2307.03172.",
    'T.B. Lee, "How long contexts hurt AI performance," *Understanding AI* '
    "(Newsletter), Nov. 2025. [Online]. Tersedia: "
    "https://www.understandingai.org/p/how-long-contexts-hurt-ai-performance",
    "P. Rajasekaran, E. Dixon, C. Ryan, dan J. Hadfield, "
    '"Effective context engineering for AI agents," *Anthropic Engineering '
    "Blog*, Sep. 2025.",
    "S.G. Patil, T. Zhang, X. Wang, dan J.E. Gonzalez, "
    '"Gorilla: Large language model connected with massive APIs," in *Proc. '
    "Adv. Neural Inf. Process. Syst. (NeurIPS)*, 2024. arXiv:2305.15334.",
    'Y. Qin et al., "ToolLLM: Facilitating large language models to master '
    '16000+ real-world APIs," in *Proc. Int. Conf. Learn. Representations '
    "(ICLR)*, 2024. arXiv:2307.16789.",
    'E. Lumer, "Toolshed: Scale tool-equipped agents with advanced RAG-tool '
    'fusion and tool knowledge bases," arXiv:2410.14594, 2024.',
    "N. Gaurav, A. Akarsh, A. Ranjan, dan M. Bajaj, "
    '"Dynamic ReAct: Scalable tool selection for large-scale MCP environments," '
    "arXiv:2509.20386, 2025.",
    'J. Jia dan Q. Li, "AutoTool: Efficient tool selection for large language '
    'model agents," in *Proc. AAAI Conf. Artif. Intell. (AAAI)*, 2026. '
    "arXiv:2511.14650.",
    "W. Wang, P. Niu, Z. Xu, et al., "
    '"MCP-Flow: Facilitating LLM agents to master real-world, diverse and '
    'scaling MCP tools," in *Proc. Assoc. Comput. Linguistics (ACL)*, 2026. '
    "arXiv:2510.24284.",
    'P. CN, "Build scalable AI agent systems using Amazon Bedrock agent '
    'registry," *AWS Machine Learning Blog*, Apr. 2026.',
    "A.R. Hevner, S.T. March, J. Park, dan S. Ram, "
    '"Design science in information systems research," *MIS Quarterly*, '
    "vol. 28, no. 1, pp. 75–105, Mar. 2004.",
    'T.B. Lee, "Context rot: The emerging challenge that could hold back LLM '
    'progress," *Understanding AI* (Newsletter), Nov. 2025. [Online]. '
    "Tersedia: https://www.understandingai.org/p/context-rot-the-emerging-challenge",
    '"Context length alone hurts LLM performance despite perfect retrieval," '
    "arXiv:2510.05381, 2025.",
    'Y. Qu et al., "Towards completeness-oriented tool retrieval for large '
    'language models," in *Proc. ACM Int. Conf. Inf. Knowl. Manag. (CIKM)*, '
    "2024. arXiv:2405.16089.",
    'S.G. Patil et al., "The Berkeley function calling leaderboard (BFCL): '
    'From tool use to agentic evaluation of large language models," in *Proc. '
    "Int. Conf. Mach. Learn. (ICML)*, 2025.",
    '"ScaleMCP: Dynamic and auto-synchronizing model context protocol tools '
    'for LLM agents," arXiv:2505.06416, 2025.',
    'TrueFoundry, "What is AI agent registry — A complete guide," *TrueFoundry '
    "Blog*, Sep. 2025. [Online]. Tersedia: "
    "https://www.truefoundry.com/blog/ai-agent-registry",
    'I. Kolchinsky, "Tool RAG: The next breakthrough in scalable AI agents," '
    "*Red Hat Emerging Technologies Blog*, Nov. 2025.",
    'Pydantic, "Pydantic AI — Toolsets," *Pydantic AI Documentation*, 2025. '
    "[Online]. Tersedia: https://ai.pydantic.dev/",
    'C. Richardson, "Pattern: API gateway / Backends for frontends," '
    "*microservices.io*, 2014–2025. [Online]. Tersedia: "
    "https://microservices.io/patterns/apigateway.html",
]

# Lampiran (placeholder) sesuai sistematika UNSIA
LAMPIRAN_ITEMS = [
    "Lampiran A: Source Code Tool Registry (cuplikan kunci)",
    "Lampiran B: Dataset 100 Query (JSON)",
    "Lampiran C: Raw Results CSV",
    "Lampiran D: Script Analisis Statistik (Python)",
    "Lampiran E: Bukti Similarity Check Turnitin (≤30%)",
]

FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)

# Pola inline markdown: bold > italic > code
_TOKEN_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
# Referensi gambar markdown: ![alt](path)
_IMG_RE = re.compile(r"^!\[.*?\]\((.+?)\)\s*$")


# --------------------------------------------------------------------------- #
# Helper XML tingkat rendah
# --------------------------------------------------------------------------- #
def _set_cell_east_asian(run):
    """Pastikan font diterapkan ke semua script (latin + east asian)."""
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), run.font.name or FONT)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _fld_simple(instr, cached="1"):
    """Bangun field Word sederhana <w:fldSimple> dengan hasil cache."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = cached
    r.append(t)
    fld.append(r)
    return fld


def set_page_number_format(section, fmt, start=None):
    sectPr = section._sectPr
    existing = sectPr.find(qn("w:pgNumType"))
    if existing is not None:
        sectPr.remove(existing)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))
    sectPr.append(pg)


def footer_page_number(section, align=WD_ALIGN_PARAGRAPH.CENTER):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.text = ""
    p.alignment = align
    add_page_number_field(p)


def add_toc_field(doc, switches='TOC \\o "1-3" \\h \\z \\u'):
    p = doc.add_paragraph()
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = switches
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Klik kanan di sini lalu pilih 'Update Field' untuk mengisi daftar."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(begin)
    r.append(instr)
    r.append(sep)
    r.append(placeholder)
    r.append(end)


def _set_borders(element, edges):
    """edges: dict edge->bool (True=single line, False=none)."""
    borders = OxmlElement("w:tblBorders") if element.tag.endswith("}tblPr") else None
    if borders is None:
        # cell border container
        borders = OxmlElement("w:tcBorders")
    for edge, on in edges.items():
        e = OxmlElement(f"w:{edge}")
        if on:
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), "8")
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), "000000")
        else:
            e.set(qn("w:val"), "none")
        borders.append(e)
    element.append(borders)


def keep_table_together(table, repeat_header=True):
    """Cegah baris terbelah antar-halaman; ulangi baris header di tiap halaman."""
    for ri, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))
        if repeat_header and ri == 0:
            trPr.append(OxmlElement("w:tblHeader"))


def three_line_table_borders(table):
    """Garis batas horizontal atas & bawah tabel saja + garis bawah header."""
    tblPr = table._tbl.tblPr
    _set_borders(
        tblPr,
        {
            "top": True,
            "bottom": True,
            "left": False,
            "right": False,
            "insideH": False,
            "insideV": False,
        },
    )
    # garis bawah baris header
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        _set_borders(
            tcPr,
            {"top": False, "bottom": True, "left": False, "right": False},
        )


# --------------------------------------------------------------------------- #
# Style dasar dokumen
# --------------------------------------------------------------------------- #
def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    for name, size, align in (
        ("Heading 1", 12, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", 12, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 12, WD_ALIGN_PARAGRAPH.LEFT),
    ):
        st = doc.styles[name]
        st.font.name = FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = BLACK
        st.font.italic = False
        st.paragraph_format.alignment = align
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    cap = doc.styles["Caption"]
    cap.font.name = FONT
    cap.font.size = Pt(12)
    cap.font.italic = False
    cap.font.bold = False
    cap.font.color.rgb = BLACK
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(6)
    cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def setup_page(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(4.0)
    section.right_margin = Cm(3.0)


# --------------------------------------------------------------------------- #
# Paragraf helper
# --------------------------------------------------------------------------- #
def add_runs(paragraph, text):
    for part in _TOKEN_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Courier New"
        else:
            paragraph.add_run(part)


def body_paragraph(doc, text, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    add_runs(p, text)
    return p


def centered(doc, text, bold=False, size=12, italic=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    add_runs(p, text)
    return p


# --------------------------------------------------------------------------- #
# Render tabel markdown
# --------------------------------------------------------------------------- #
def add_caption(doc, kind, chap_no, text, gambar=False, keep_with_next=False):
    """Caption ber-SEQ otomatis: 'Tabel 2.<SEQ>' / 'Gambar 2.<SEQ>'.

    `\\s 1` mereset nomor SEQ di tiap Heading 1 (per bab); prefiks nomor bab
    ditulis literal. Field ini yang dikoleksi Daftar Tabel/Gambar (TOC \\c).
    """
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if keep_with_next:
        p.paragraph_format.keep_with_next = True
    p.add_run(f"{kind} {chap_no}.")
    p._p.append(_fld_simple(f" SEQ {kind} \\* ARABIC \\s 1 "))
    txt = text.strip()
    if gambar and not txt.endswith("."):
        txt += "."
    p.add_run(f" {txt}")
    return p


def add_image_centered(doc, rel_path, max_w_cm=12.0, max_h_cm=18.0):
    """Sisipkan gambar di tengah, diskalakan agar tidak melebihi area halaman."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    path = rel_path if rel_path.is_absolute() else (ROOT / rel_path)
    if not path.exists():
        r = p.add_run(f"[Gambar tidak ditemukan: {rel_path}]")
        r.italic = True
        return p
    w, h = Image.open(path).size
    width_cm = max_w_cm
    if (h / w) * width_cm > max_h_cm:
        width_cm = max_h_cm * w / h
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    return p


def _set_para_shading(p, fill="F2F2F2"):
    """Beri warna latar (shading) pada paragraf via w:shd di pPr."""
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def add_code_block(doc, code_lines):
    """Render blok kode (fenced ```), monospace single-spacing dengan shading.

    Tidak terbelah antar-halaman bila ≤ 30 baris; selebihnya dibiarkan mengalir.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.left_indent = Cm(0.5)
    pf.right_indent = Cm(0.2)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.keep_together = len(code_lines) <= 30
    _set_para_shading(p, "F2F2F2")
    for idx, raw in enumerate(code_lines):
        # ganti tab dengan 4 spasi, pertahankan indentasi awal baris
        text = raw.replace("\t", "    ")
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        _set_cell_east_asian(run)
        if idx < len(code_lines) - 1:
            run.add_break()
    return p


def render_table(doc, rows):
    """rows: list of list[str] sudah dipisah kolom (termasuk header)."""
    ncol = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncol)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.paragraphs[0].text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_after = Pt(0)
            add_runs(p, cell_text)
            for run in p.runs:
                run.font.size = Pt(10)
                if i == 0:
                    run.bold = True
    three_line_table_borders(table)
    keep_table_together(table, repeat_header=True)
    return table


# --------------------------------------------------------------------------- #
# Parser markdown bab
# --------------------------------------------------------------------------- #
_TABLE_ROW_RE = re.compile(r"^\s*\|(.+)\|\s*$")
_SEP_ROW_RE = re.compile(r"^\s*\|?[\s:|-]+\|?\s*$")
_CAPTION_RE = re.compile(r"^(Tabel|Gambar)\s+\d", re.IGNORECASE)
_CODE_CAPTION_RE = re.compile(r"^Kode Program\s+\d+\.\d+\s", re.IGNORECASE)
_FENCE_RE = re.compile(r"^```")


def split_table_row(line):
    inner = _TABLE_ROW_RE.match(line).group(1)
    return [c.strip() for c in inner.split("|")]


def _next_nonblank_is_table(lines, idx):
    """True bila baris non-kosong berikutnya (mulai idx) adalah baris tabel."""
    while idx < len(lines):
        if not lines[idx].strip():
            idx += 1
            continue
        return bool(_TABLE_ROW_RE.match(lines[idx]))
    return False


def render_markdown(doc, md_text, captions, chap_no):
    lines = md_text.splitlines()
    # buang baris judul bab (## diawali '# ') di awal file — ditangani terpisah
    i = 0
    n = len(lines)
    last_was_image = False  # untuk deteksi caption Gambar (caption di bawah gambar)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # tandai apakah baris sebelumnya (non-kosong) adalah gambar, lalu reset
        was_image = last_was_image
        last_was_image = False

        # judul bab (sudah dirender manual) — lewati H1 di paling atas
        if stripped.startswith("# "):
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        # blok kode berpagar ```
        if _FENCE_RE.match(stripped):
            i += 1
            code_lines: list[str] = []
            while i < n and not _FENCE_RE.match(lines[i].strip()):
                code_lines.append(lines[i])
                i += 1
            i += 1  # lewati pagar penutup
            # buang baris kosong di ujung blok
            while code_lines and not code_lines[0].strip():
                code_lines.pop(0)
            while code_lines and not code_lines[-1].strip():
                code_lines.pop()
            if code_lines:
                add_code_block(doc, code_lines)
            continue

        # caption listing kode: "Kode Program X.Y ..." (manual, tanpa SEQ)
        if _CODE_CAPTION_RE.match(stripped):
            p = doc.add_paragraph(style="Caption")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            add_runs(p, stripped)
            i += 1
            continue

        # tabel
        if _TABLE_ROW_RE.match(line):
            block = []
            while i < n and _TABLE_ROW_RE.match(lines[i]):
                block.append(lines[i])
                i += 1
            rows = [
                split_table_row(b) for b in block if not _SEP_ROW_RE.match(b)
            ]
            if rows:
                render_table(doc, rows)
            continue

        # gambar: ![alt](path) — disisipkan di tengah, fit halaman
        m_img = _IMG_RE.match(stripped)
        if m_img:
            add_image_centered(doc, Path(m_img.group(1).strip()))
            last_was_image = True
            i += 1
            continue

        # caption tabel/gambar via SEQ field (auto-numbering + Daftar otomatis).
        # Deteksi sadar-konteks agar kalimat yang kebetulan diawali "Tabel X.Y ..."
        # tidak ikut dianggap caption: caption Tabel harus diikuti tabel, caption
        # Gambar harus tepat mengikuti sebuah gambar.
        if _CAPTION_RE.match(stripped):
            m = re.match(
                r"^(Tabel|Gambar)\s+\d+\.\d+\s+(.*)$", stripped, re.IGNORECASE
            )
            if m:
                kind = m.group(1).capitalize()
                is_gambar = kind.lower() == "gambar"
                is_caption = (
                    (is_gambar and was_image)
                    or (not is_gambar and _next_nonblank_is_table(lines, i + 1))
                )
                if is_caption:
                    # Tabel: caption di ATAS tabel → glue ke tabel berikutnya.
                    # Gambar: caption di BAWAH gambar → tidak perlu keep_with_next.
                    add_caption(
                        doc,
                        kind,
                        chap_no,
                        m.group(2),
                        gambar=is_gambar,
                        keep_with_next=not is_gambar,
                    )
                    captions.append(stripped)
                    i += 1
                    continue
            # bukan caption sejati → biarkan jatuh ke paragraf biasa di bawah

        # heading sub-bab
        if stripped.startswith("### "):
            heading(doc, stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            heading(doc, stripped[3:].strip(), level=2)
            i += 1
            continue

        # subjudul "Referensi Bab ..." (baris full-bold) — jangan jadi yatim
        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(12)
            r = p.add_run(stripped[2:-2])
            r.bold = True
            i += 1
            continue

        # daftar bernomor / berhuruf
        if re.match(r"^\d+\.\s+", stripped) or re.match(r"^[a-z]\)\s+", stripped):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            add_runs(p, stripped)
            i += 1
            continue

        # referensi IEEE [n] — rata kiri, hanging indent, tiap entri utuh
        if re.match(r"^\[\d+\]\s", stripped):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-1.0)
            p.paragraph_format.keep_together = True
            p.paragraph_format.space_after = Pt(6)
            add_runs(p, stripped)
            i += 1
            continue

        # paragraf biasa
        body_paragraph(doc, stripped)
        i += 1


# --------------------------------------------------------------------------- #
# Front matter
# --------------------------------------------------------------------------- #
def build_cover(doc):
    centered(doc, "TUGAS AKHIR", bold=True, size=16)
    centered(doc, "", size=16)
    for ln in _wrap_title(META["judul"]):
        centered(doc, ln, bold=True, size=16)
    for _ in range(3):
        centered(doc, "")
    centered(
        doc,
        f"Sebagai Salah Satu Syarat untuk Memperoleh Gelar {META['gelar']} pada",
        size=14,
    )
    centered(doc, f"Program Studi {META['prodi']}", size=14)
    centered(doc, "")
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Cm(3.5))
    centered(doc, "")
    centered(doc, "Oleh :", bold=True, size=12)
    centered(doc, f"{META['nim']}, {META['nama']}", bold=True, size=12)
    for _ in range(2):
        centered(doc, "")
    centered(doc, f"PROGRAM STUDI PJJ {META['prodi'].upper()}", bold=True, size=14)
    centered(doc, "UNIVERSITAS SIBER ASIA", bold=True, size=14)
    centered(doc, f"{META['bulan_tahun']}", bold=True, size=14)


def _wrap_title(title, width=48):
    words = title.split()
    lines, cur = [], ""
    for w in words:
        if len(cur) + len(w) + 1 > width:
            lines.append(cur)
            cur = w
        else:
            cur = f"{cur} {w}".strip()
    if cur:
        lines.append(cur)
    return lines


def build_orisinalitas(doc):
    heading(doc, "HALAMAN PERNYATAAN ORISINALITAS", level=1)
    centered(doc, "")
    body_paragraph(doc, "Yang bertanda tangan di bawah ini:", indent=False)
    for label, val in (
        ("Nama", META["nama"]),
        ("NIM", META["nim"]),
        ("Judul Tugas Akhir", META["judul"]),
    ):
        p = doc.add_paragraph()
        p.add_run(f"{label}").bold = False
        p.add_run("\t: ")
        add_runs(p, val)
    centered(doc, "")
    body_paragraph(
        doc,
        "Menyatakan dengan sesungguhnya bahwa Tugas Akhir ini merupakan hasil "
        "penelitian, pemikiran, dan pemaparan asli penulis sendiri. Penulis tidak "
        "mencantumkan tanpa pengakuan bahan-bahan yang telah dipublikasikan "
        "sebelumnya atau ditulis oleh orang lain, atau sebagai bahan yang pernah "
        "diajukan untuk gelar atau ijazah pada Universitas Siber Asia atau "
        "Perguruan Tinggi lainnya.",
    )
    body_paragraph(
        doc,
        "Apabila di kemudian hari terdapat penyimpangan dan ketidakbenaran dalam "
        "pernyataan ini, maka penulis bersedia menerima sanksi akademik sesuai "
        "dengan peraturan yang berlaku di Universitas Siber Asia.",
    )
    body_paragraph(doc, "Demikian pernyataan ini penulis buat.", indent=False)
    centered(doc, "")
    body_paragraph(doc, f"{META['kota']}, ___/___/______", indent=False)
    body_paragraph(doc, "Yang membuat pernyataan,", indent=False)
    for _ in range(2):
        body_paragraph(doc, "", indent=False)
    body_paragraph(doc, "(materai Rp10.000,-)", indent=False)
    centered(doc, "")
    body_paragraph(doc, META["nama"], indent=False)


def build_pengesahan(doc):
    heading(doc, "HALAMAN PENGESAHAN", level=1)
    centered(doc, "")
    body_paragraph(doc, "Tugas Akhir ini diajukan oleh:", indent=False)
    for label, val in (
        ("Nama", META["nama"]),
        ("NIM", META["nim"]),
        ("Program Studi", META["prodi"]),
        ("Judul Tugas Akhir", META["judul"]),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.add_run(label)
        p.add_run("\t: ")
        add_runs(p, val)
    body_paragraph(
        doc,
        f"Telah berhasil dipertahankan di hadapan Dewan Penguji dan diterima "
        f"sebagai bagian persyaratan yang diperlukan untuk memperoleh gelar "
        f"{META['gelar']} pada Program Studi {META['prodi']} Universitas Siber Asia.",
    )
    centered(doc, "")
    _signature_block(
        doc,
        [
            ("Ketua Sidang / Dosen Penguji I", "…………………………"),
            ("Dosen Penguji II", "…………………………"),
        ],
    )
    _signature_block(
        doc,
        [
            ("Dosen Pembimbing I", META["pembimbing"]),
            ("Dosen Pembimbing II", "…………………………"),
        ],
    )
    centered(doc, "Ketua Program Studi", space_after=36)
    centered(doc, "…………………………")
    centered(doc, "")
    body_paragraph(doc, f"Ditetapkan di : {META['kota']}", indent=False)
    body_paragraph(doc, "Tanggal         : _____/_____/202_", indent=False)


def _signature_block(doc, pairs):
    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (role, _name) in enumerate(pairs):
        c = t.cell(0, j).paragraphs[0]
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.add_run(role)
    for j, (_role, name) in enumerate(pairs):
        c = t.cell(1, j).paragraphs[0]
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_before = Pt(36)
        add_runs(c, name)
    keep_table_together(t, repeat_header=False)
    centered(doc, "")


def build_kata_pengantar(doc):
    heading(doc, "KATA PENGANTAR", level=1)
    centered(doc, "")
    body_paragraph(
        doc,
        "Puji syukur penulis panjatkan kepada Tuhan Yang Maha Esa, karena atas "
        "berkat dan rahmat-Nya penulis dapat menyelesaikan Tugas Akhir ini. "
        f"Penulisan Tugas Akhir ini dilakukan dalam rangka memenuhi salah satu "
        f"syarat untuk mencapai gelar {META['gelar']} pada Program Studi "
        f"{META['prodi']} Universitas Siber Asia.",
    )
    body_paragraph(
        doc,
        "Penulis menyadari bahwa tanpa bantuan dan bimbingan dari berbagai pihak, "
        "dari masa perkuliahan sampai pada penyusunan Tugas Akhir ini, sangatlah "
        "sulit bagi penulis untuk menyelesaikan Tugas Akhir ini. Oleh karena itu, "
        "penulis mengucapkan terima kasih kepada:",
    )
    for item in (
        f"{META['pembimbing']}, selaku dosen pembimbing yang telah menyediakan "
        "waktu, tenaga, dan pikiran untuk mengarahkan penulis dalam penyusunan "
        "Tugas Akhir ini;",
        "tim dan komunitas pengembang platform zerlo.id yang telah banyak "
        "membantu dalam penyediaan data dan lingkungan eksperimen;",
        "orang tua dan keluarga penulis yang telah memberikan bantuan dukungan "
        "material dan moral; dan",
        "sahabat yang telah banyak membantu penulis dalam menyelesaikan Tugas "
        "Akhir ini.",
    ):
        p = doc.add_paragraph(style="List Number")
        add_runs(p, item)
    body_paragraph(
        doc,
        "Akhir kata, penulis berharap Tuhan Yang Maha Esa berkenan membalas segala "
        "kebaikan semua pihak yang telah membantu. Semoga Tugas Akhir ini membawa "
        "manfaat bagi pengembangan ilmu.",
    )
    centered(doc, "")
    body_paragraph(doc, f"{META['kota']}, {META['bulan_tahun']}", indent=False)
    body_paragraph(doc, "Penulis", indent=False)


def build_abstrak(doc):
    """Abstrak dwibahasa dari `bab/abstrak.md` (spasi 1, justify, satu paragraf
    per bahasa diikuti baris Kata Kunci/Keywords)."""
    md = (BAB_DIR / "abstrak.md").read_text(encoding="utf-8")
    blocks = [b.strip() for b in re.split(r"^# ", md, flags=re.M) if b.strip()]
    for n, blk in enumerate(blocks):
        lines = [ln for ln in blk.split("\n") if ln.strip()]
        title = lines[0].strip()
        if n:
            centered(doc, "")  # pemisah antar-bahasa
        heading(doc, title, level=1)
        for raw in lines[1:]:
            line = raw.strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            if not (line.startswith("**Kata Kunci") or line.startswith("**Keywords")):
                p.paragraph_format.first_line_indent = Cm(1.27)
            add_runs(p, line)


def build_daftar(doc, title, toc_switches, note):
    heading(doc, title, level=1)
    centered(doc, "")
    add_toc_field(doc, toc_switches)
    if note:
        p = doc.add_paragraph()
        r = p.add_run(note)
        r.italic = True
        r.font.size = Pt(10)


# --------------------------------------------------------------------------- #
# Bab body & placeholder
# --------------------------------------------------------------------------- #
def build_chapter(doc, entry, captions, chap_no, page_break=True):
    md_file, label, title, outline = entry
    if page_break:
        doc.add_page_break()
    heading(doc, f"{label} {title}", level=1)
    if md_file:
        md_text = (BAB_DIR / md_file).read_text(encoding="utf-8")
        render_markdown(doc, md_text, captions, chap_no)
    else:
        p = doc.add_paragraph()
        r = p.add_run(f"[{label} BELUM DITULIS]")
        r.bold = True
        r.italic = True
        body_paragraph(
            doc,
            "Rencana sub-bab (mengikuti sistematika UNSIA untuk skema Tugas Akhir "
            "Prototype):",
            indent=False,
        )
        for sub in outline or []:
            heading(doc, sub, level=2)
            ph = doc.add_paragraph()
            ph.paragraph_format.first_line_indent = Cm(1.27)
            rr = ph.add_run("[konten belum ditulis]")
            rr.italic = True


def build_daftar_pustaka(doc):
    doc.add_page_break()
    heading(doc, "DAFTAR PUSTAKA", level=1)
    centered(doc, "")
    for n, ref in enumerate(REFERENCES, start=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(6)
        pf.left_indent = Cm(1.0)
        pf.first_line_indent = Cm(-1.0)  # hanging indent ala IEEE
        p.add_run(f"[{n}] ")
        add_runs(p, ref)


def _lampiran_heading(doc, title):
    h = doc.add_paragraph(style="Heading 2")
    h.add_run(title)


def _embed_source(doc, rel_path, intro, max_lines=0):
    """Sisipkan isi file sumber sebagai code block (dibaca saat build agar
    selalu sinkron dengan repositori)."""
    body_paragraph(doc, intro, indent=False)
    text = (ROOT / rel_path).read_text(encoding="utf-8")
    lines = text.splitlines()
    truncated = False
    if max_lines and len(lines) > max_lines:
        lines = lines[:max_lines]
        truncated = True
    centered(doc, rel_path, italic=True, size=10)
    add_code_block(doc, lines)
    if truncated:
        body_paragraph(
            doc,
            f"(Cuplikan {max_lines} baris pertama; berkas lengkap tersedia pada "
            f"repositori penelitian: {rel_path}.)",
            indent=False,
        )


def build_lampiran(doc):
    doc.add_page_break()
    heading(doc, "LAMPIRAN", level=1)

    # --- Lampiran A: Source code Tool Registry ---
    _lampiran_heading(doc, "Lampiran A: Source Code Tool Registry (Cuplikan Kunci)")
    _embed_source(
        doc,
        "src/tool_registry_eval/registry.py",
        "Modul registry berisi struktur metadata dan fungsi penyaringan "
        "`registry_filter()` yang menjadi artefak utama penelitian ini.",
    )

    # --- Lampiran B: Dataset 100 query ---
    _lampiran_heading(doc, "Lampiran B: Dataset 100 Kueri Evaluasi (JSON)")
    body_paragraph(
        doc,
        "Dataset evaluasi terdiri dari 100 kueri berbahasa Indonesia (50 "
        "single-domain, 30 cross-domain, 20 adversarial). Berikut cuplikan "
        "representatif; berkas lengkap tersedia pada lampiran/dataset-100-query.json.",
        indent=False,
    )
    try:
        import json as _json

        data = _json.loads(
            (ROOT / "lampiran" / "dataset-100-query.json").read_text(encoding="utf-8")
        )
        sample, seen = [], {}
        for q in data:
            t = q.get("query_type")
            if seen.get(t, 0) < 2:
                sample.append(q)
                seen[t] = seen.get(t, 0) + 1
        snippet = _json.dumps(sample, ensure_ascii=False, indent=2)
        add_code_block(doc, snippet.splitlines())
    except FileNotFoundError:
        body_paragraph(doc, "[dataset-100-query.json belum di-generate]", indent=False)

    # --- Lampiran C: Raw results CSV ---
    _lampiran_heading(doc, "Lampiran C: Hasil Mentah Eksperimen (CSV)")
    _embed_source(
        doc,
        "outputs/gemini-native-v2/summary.csv",
        "Ringkasan agregat per skenario dan mode (Gemini Native v2, n=558), "
        "diproduksi otomatis dari rekaman JSONL inkremental.",
    )
    _embed_source(
        doc,
        "outputs/gemini-native-v2/statistical_tests.csv",
        "Hasil uji statistik penghematan token (Wilcoxon, Cohen's d, CI 95%) dan "
        "akurasi per skenario.",
    )

    # --- Lampiran D: Script analisis statistik ---
    _lampiran_heading(doc, "Lampiran D: Script Analisis Statistik (Python)")
    _embed_source(
        doc,
        "src/tool_registry_eval/measure.py",
        "Modul pengukuran yang menghitung uji Wilcoxon signed-rank berpasangan, "
        "ukuran efek Cohen's d, dan selang kepercayaan 95% atas penghematan token.",
    )

    # --- Lampiran E: Bukti Turnitin ---
    _lampiran_heading(doc, "Lampiran E: Bukti Similarity Check Turnitin (≤30%)")
    body_paragraph(
        doc,
        "Pemeriksaan kemiripan (*similarity check*) dilakukan menggunakan Turnitin "
        "atas naskah final Bab I–V. Bukti berupa halaman ringkasan Turnitin "
        "(*similarity report*) dengan skor kemiripan maksimal 30% dilampirkan pada "
        "versi cetak setelah naskah dinyatakan final oleh dosen pembimbing.",
        indent=False,
    )
    ph = doc.add_paragraph()
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ph.add_run(
        "[Lampiran bukti Turnitin disisipkan pada tahap finalisasi naskah]"
    ).italic = True


# --------------------------------------------------------------------------- #
# Main
# --------------------------------------------------------------------------- #
def main():
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    setup_styles(doc)

    # --- Section 1: COVER (tanpa nomor halaman) ---
    sec_cover = doc.sections[0]
    setup_page(sec_cover)
    build_cover(doc)

    # --- Section 2: FRONT MATTER (angka Romawi kecil) ---
    sec_front = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_page(sec_front)
    set_page_number_format(sec_front, "lowerRoman", start=2)
    footer_page_number(sec_front, WD_ALIGN_PARAGRAPH.CENTER)

    build_orisinalitas(doc)
    doc.add_page_break()
    build_pengesahan(doc)
    doc.add_page_break()
    build_abstrak(doc)
    doc.add_page_break()
    build_kata_pengantar(doc)
    doc.add_page_break()
    build_daftar(
        doc,
        "DAFTAR ISI",
        'TOC \\o "1-3" \\h \\z \\u',
        "Daftar Isi otomatis: di Word tekan Ctrl+A lalu F9 untuk memperbarui.",
    )
    doc.add_page_break()
    build_daftar(
        doc,
        "DAFTAR TABEL",
        'TOC \\h \\z \\c "Tabel"',
        "Terisi otomatis dari caption ber-SEQ; di Word tekan Ctrl+A lalu F9.",
    )
    doc.add_page_break()
    build_daftar(
        doc,
        "DAFTAR GAMBAR",
        'TOC \\h \\z \\c "Gambar"',
        "Terisi otomatis dari caption ber-SEQ; di Word tekan Ctrl+A lalu F9.",
    )

    # --- Section 3: BODY (angka Arab) ---
    sec_body = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_page(sec_body)
    set_page_number_format(sec_body, "decimal", start=1)
    footer_page_number(sec_body, WD_ALIGN_PARAGRAPH.CENTER)

    captions = []
    for idx, entry in enumerate(CHAPTERS):
        # bab pertama tanpa page-break ganda setelah section break
        build_chapter(doc, entry, captions, chap_no=idx + 1, page_break=(idx != 0))

    build_daftar_pustaka(doc)
    build_lampiran(doc)

    doc.save(OUT_FILE)
    print(f"✓ Naskah tersimpan: {OUT_FILE.relative_to(ROOT)}")
    print(f"  Bab termuat   : {sum(1 for c in CHAPTERS if c[0])} / {len(CHAPTERS)}")
    print(f"  Caption tabel : {len(captions)} ({', '.join(captions) or '—'})")
    print("  Buka di Word → Ctrl+A → F9 untuk memperbarui Daftar Isi.")


if __name__ == "__main__":
    main()
